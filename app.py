import streamlit as st
import os
import google.generativeai as genai
from dotenv import load_dotenv
import tempfile
from docx import Document
from docx.shared import Pt
from io import BytesIO
import re  # <--- 新增這個套件，用來進行強力搜尋

# --- 1. 頁面設定 ---
st.set_page_config(page_title="AI 超級會議秘書", page_icon="🚀", layout="wide")

# --- 2. API 設定 ---
if "GOOGLE_API_KEY" in st.secrets:
    api_key = st.secrets["GOOGLE_API_KEY"]
else:
    load_dotenv()
    api_key = os.getenv("GOOGLE_API_KEY")

with st.sidebar:
    st.header("⚙️ 設定")
    if not api_key:
        api_key = st.text_input("Google API Key", type="password")
    st.info("💡 修正版：\n強化標籤抓取邏輯 (Regex)\n修復心智圖與 Email 空白問題")

if not api_key:
    st.warning("請設定 API Key")
    st.stop()

# --- 3. 初始化 Gemini ---
genai.configure(api_key=api_key)
# 建議維持使用 2.0-flash-001，若之後又出現問題可改回 'gemini-2.0-flash'
model = genai.GenerativeModel('gemini-2.0-flash-001')

# --- 4. Word 轉檔函式 ---
def generate_docx(content):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Microsoft JhengHei'
    style.font.size = Pt(12)
    
    for line in content.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line[2:], 0)
        elif line.startswith('## '): doc.add_heading(line[3:], 1)
        elif line.startswith('### '): doc.add_heading(line[4:], 2)
        elif line.startswith('- ') or line.startswith('* '): 
            doc.add_paragraph(line[2:], style='List Bullet')
        else: doc.add_paragraph(line)
            
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 5. 強力標籤提取函式 (關鍵修正) ---
def extract_tag_content(text, tag_name):
    """
    使用正規表達式 (Regex) 來抓取 <tag>...</tag> 之間的內容。
    re.DOTALL 讓 . 可以匹配換行符號，確保跨行內容也能抓到。
    """
    pattern = rf"<{tag_name}>(.*?)</{tag_name}>"
    match = re.search(pattern, text, re.DOTALL)
    if match:
        return match.group(1).strip()
    return "" # 抓不到就回傳空字串

# --- 6. 主介面 ---
st.title("🚀 AI 超級會議秘書 (修正版)")
st.caption("全能版：錄音/上傳 + 逐字稿 + Word + 心智圖 + Email 草稿")

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("1. 輸入音訊")
    tab1, tab2 = st.tabs(["🎙️ 錄音", "📂 上傳"])
    audio_source = None
    source_name = ""
    
    with tab1:
        mic = st.audio_input("開始錄音")
        if mic: 
            audio_source = mic
            source_name = "rec.wav"
            
    with tab2:
        up = st.file_uploader("上傳 MP3/M4A", type=["mp3","wav","m4a"])
        if up: 
            audio_source = up
            source_name = up.name

# --- 7. 核心處理 ---
if audio_source:
    with col1:
        st.success("音訊就緒")
        start_btn = st.button("🚀 開始全能分析", type="primary")

    if start_btn:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(source_name)[1] or ".wav") as tmp:
            tmp.write(audio_source.getvalue())
            tmp_path = tmp.name

        try:
            with st.spinner("🧠 AI 正在大腦風暴中 (分析、畫圖、寫信)..."):
                # A. 上傳
                upload_file = genai.upload_file(tmp_path, mime_type=audio_source.type)
                
                # B. Prompt (加強指令，告訴 AI 不要亂加 markdown 代碼塊)
                prompt = """
                你是一位頂級會議秘書。請針對這段錄音完成以下任務。
                
                【絕對規則】
                1. 請務必使用 XML 標籤 (`<tag>...</tag>`) 將不同區塊分開。
                2. **不要** 在 XML 標籤外面包覆 Markdown 代碼符號 (如 ```xml ... ```)，直接輸出標籤即可。
                3. 使用繁體中文 (台灣)。

                任務 1：逐字稿
                請包在 <transcript> ... </transcript> 標籤中。
                區分講者 (如 [講者1])，講者間需換行。

                任務 2：會議紀要 (Markdown 格式)
                請包在 <summary> ... </summary> 標籤中。
                格式：
                # 會議紀錄
                ## 🎯 主旨
                ## 🔑 決策
                ## ✅ 待辦事項 (列點)

                任務 3：心智圖語法 (Mermaid)
                請包在 <mindmap> ... </mindmap> 標籤中。
                只要輸出 `graph TD` 開頭的語法內容，不要加 ```mermaid 符號。
                結構：會議主題 -> 關鍵議題 -> 細項。

                任務 4：Email 草稿
                請包在 <email> ... </email> 標籤中。
                撰寫一封給「所有與會者」的跟進 Email。
                """
                
                response = model.generate_content([prompt, upload_file])
                full_text = response.text
                
                # C. 使用強力函式解析
                transcript = extract_tag_content(full_text, "transcript")
                summary = extract_tag_content(full_text, "summary")
                mindmap_code = extract_tag_content(full_text, "mindmap")
                email_draft = extract_tag_content(full_text, "email")

                # 如果心智圖還是包含 markdown 符號，手動清理
                if mindmap_code:
                    mindmap_code = mindmap_code.replace("```mermaid", "").replace("```", "").strip()

                # D. 顯示結果
                with col2:
                    st.divider()
                    
                    # 預防機制：如果真的全部抓失敗，顯示警告
                    if not transcript and not summary:
                        st.error("⚠️ 解析失敗，AI 可能沒有依照格式輸出。請查看下方的「原始回傳內容」。")
                    
                    res_tab1, res_tab2, res_tab3, res_tab4 = st.tabs(["📊 心智圖", "📝 正式報告", "✉️ Email 草稿", "💬 逐字稿"])
                    
                    with res_tab1:
                        st.subheader("會議結構可視化")
                        if mindmap_code:
                            st.mermaid(mindmap_code)
                        else:
                            st.info("本次分析未能生成心智圖結構。")

                    with res_tab2:
                        st.subheader("會議紀要")
                        if summary:
                            st.markdown(summary)
                            # 合併下載
                            full_doc = summary + "\n\n---\n\n" + (transcript if transcript else "")
                            docx = generate_docx(full_doc)
                            st.download_button("📥 下載 Word 報告", docx, "minutes.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        else:
                            st.warning("無摘要內容")

                    with res_tab3:
                        st.subheader("跟進郵件草稿")
                        if email_draft:
                            st.text_area("複製內容：", email_draft, height=300)
                        else:
                            st.info("無 Email 草稿")

                    with res_tab4:
                        st.subheader("完整對話")
                        st.text_area("逐字稿", transcript if transcript else "無內容", height=400)
                    
                    # E. 除錯區塊 (關鍵！如果又失敗，點開這裡看真相)
                    with st.expander("🛠️ 開發者除錯模式 (查看 AI 原始回傳)"):
                        st.text(full_text)

        except Exception as e:
            st.error(f"錯誤: {e}")
        finally:
            if os.path.exists(tmp_path): os.remove(tmp_path)
